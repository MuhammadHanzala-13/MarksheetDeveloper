import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from PIL import Image
import os
import json

def calculate_grades(df):
    """Calculates total marks, GPA, and grades."""
    for index, row in df.iterrows():
        mid_term_marks = json.loads(row['Mid-term Marks'].replace("'", '"'))
        final_marks = json.loads(row['Final Marks'].replace("'", '"'))
        subjects = row['Subject'].split(',')

        total_marks_dict = {}
        for subject in subjects:
            mid_mark = mid_term_marks.get(subject, 0)
            final_mark = final_marks.get(subject, 0)
            total_marks_dict[subject] = 0.3 * mid_mark + 0.7 * final_mark
        
        df.at[index, 'Total Marks'] = json.dumps(total_marks_dict)
        
        total_percentage = sum(total_marks_dict.values()) / len(subjects)
        df.at[index, 'Percentage'] = total_percentage
        
        gpa = (total_percentage / 100) * 4.0
        df.at[index, 'GPA'] = f"{gpa:.2f}"

        if total_percentage >= 90:
            grade = 'A+'
        elif 80 <= total_percentage < 90:
            grade = 'A'
        elif 70 <= total_percentage < 80:
            grade = 'B'
        elif 60 <= total_percentage < 70:
            grade = 'C'
        elif 50 <= total_percentage < 60:
            grade = 'D'
        else:
            grade = 'F'
        df.at[index, 'Grade'] = grade
        
    return df

def generate_word_marksheet(student_data, school_name, school_logo_path):
    """Generates a Word document for a single student's marksheet."""
    document = Document()
    
    # Add header with school logo and name
    if os.path.exists(school_logo_path):
        header_section = document.sections[0].header
        header_paragraph = header_section.paragraphs[0]
        logo_run = header_paragraph.add_run()
        logo_run.add_picture(school_logo_path, width=Inches(1.0))
        header_paragraph.add_run(f'\t{school_name}').bold = True
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print(f"Warning: School logo not found at {school_logo_path}")
        document.add_heading(school_name, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('Marksheet', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Student Profile section
    profile_table = document.add_table(rows=1, cols=2)
    profile_cell = profile_table.cell(0, 0)
    profile_cell.text = f"Name: {student_data['Student Name']}\n" \
                       f"Roll No: {student_data['Roll No']}\n" \
                       f"Department: {student_data['Department']}"

    photo_cell = profile_table.cell(0, 1)
    if os.path.exists(student_data['Photo Path']):
        photo_run = photo_cell.paragraphs[0].add_run()
        photo_run.add_picture(student_data['Photo Path'], width=Inches(1.25))
    else:
        print(f"Warning: Student photo not found at {student_data['Photo Path']}")
        photo_cell.text = "Photo not available"

    # Marks table
    document.add_heading('Marks Details', level=3)
    subjects = student_data['Subject'].split(',')
    marks_table = document.add_table(rows=1, cols=4)
    marks_table.style = 'Table Grid'
    hdr_cells = marks_table.rows[0].cells
    hdr_cells[0].text = 'Subject'
    hdr_cells[1].text = 'Mid-term Marks'
    hdr_cells[2].text = 'Final Marks'
    hdr_cells[3].text = 'Total Marks'

    mid_term_marks = json.loads(student_data['Mid-term Marks'].replace("'", '"'))
    final_marks = json.loads(student_data['Final Marks'].replace("'", '"'))
    total_marks = json.loads(student_data['Total Marks'])

    for subject in subjects:
        row_cells = marks_table.add_row().cells
        row_cells[0].text = subject
        row_cells[1].text = str(mid_term_marks.get(subject, 'N/A'))
        row_cells[2].text = str(final_marks.get(subject, 'N/A'))
        row_cells[3].text = f"{total_marks.get(subject, 0):.2f}"

    # Summary
    document.add_heading('Summary', level=3)
    document.add_paragraph(f"Percentage: {student_data['Percentage']:.2f}%\n"
                           f"GPA: {student_data['GPA']}\n"
                           f"Grade: {student_data['Grade']}")

    # Save the document
    output_filename = f"output/{student_data['Student Name']}_marksheet.docx"
    document.save(output_filename)
    print(f"Generated marksheet for {student_data['Student Name']} at {output_filename}")

def generate_excel_master_sheet(df, school_name):
    """Generates an Excel master result sheet."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Master Result Sheet"

    ws.append([school_name])
    ws.merge_cells('A1:G1')
    ws['A1'].font = Font(bold=True, size=16)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    # Add headers
    headers = ['Student Name', 'Roll No', 'Department', 'Percentage', 'GPA', 'Grade']
    ws.append(headers)

    # Style headers
    for cell in ws[2]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')

    # Adjust column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 10
    ws.column_dimensions['F'].width = 10

    for r_idx, row in df.iterrows():
        row_data = [row['Student Name'], row['Roll No'], row['Department'], 
                    f"{row['Percentage']:.2f}%", row['GPA'], row['Grade']]
        ws.append(row_data)
        # Center align the data in all cells of the new row
        for cell in ws[ws.max_row]:
            cell.alignment = Alignment(horizontal='center')
            if row['Grade'] == 'F':
                cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")

    output_filename = "output/master_result_sheet.xlsx"
    wb.save(output_filename)
    print(f"Generated master result sheet at {output_filename}")

def main():
    """Main function to drive the marksheet generation."""
    # Configuration
    school_name = "GitHub Copilot School of AI"
    school_logo_path = "assets/logos/logo.png"
    data_path = "data/marks_data.csv"

    # Create dummy assets if they don't exist
    if not os.path.exists(school_logo_path):
        try:
            img = Image.new('RGB', (100, 100), color = 'red')
            img.save(school_logo_path)
            print(f"Created dummy logo at {school_logo_path}")
        except Exception as e:
            print(f"Could not create dummy logo: {e}")


    # Load data
    try:
        df = pd.read_csv(data_path)
    except FileNotFoundError:
        print(f"Error: Data file not found at {data_path}")
        return

    # Create dummy student photos if they don't exist
    for photo_path in df['Photo Path']:
        if not os.path.exists(photo_path):
            try:
                img = Image.new('RGB', (100, 120), color = 'blue')
                os.makedirs(os.path.dirname(photo_path), exist_ok=True)
                img.save(photo_path)
                print(f"Created dummy photo at {photo_path}")
            except Exception as e:
                print(f"Could not create dummy photo {photo_path}: {e}")


    # Calculations
    df = calculate_grades(df)

    # Generate documents
    for index, row in df.iterrows():
        generate_word_marksheet(row, school_name, school_logo_path)
    
    generate_excel_master_sheet(df, school_name)

if __name__ == "__main__":
    main()


    