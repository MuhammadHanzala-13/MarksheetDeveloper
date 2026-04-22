import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import io
import base64
import matplotlib.pyplot as plt
import seaborn as sns

# --- Core Marksheet Generation Logic ---

def calculate_single_student_grades(student_info, mid_term_weight, final_term_weight, passing_percentage):
    """Calculates total marks, GPA, and grades for a single student from form data."""
    total_marks_dict = {}
    subjects = list(student_info['marks'].keys())

    for subject, marks in student_info['marks'].items():
        mid_mark = marks['Mid-term']
        final_mark = marks['Final']
        total_marks_dict[subject] = (mid_term_weight * mid_mark) + (final_term_weight * final_mark)

    student_info['total_marks'] = total_marks_dict
    
    total_percentage = sum(total_marks_dict.values()) / len(subjects) if subjects else 0
    student_info['percentage'] = total_percentage

    gpa = (total_percentage / 100) * 4.0
    student_info['gpa'] = f"{gpa:.2f}"

    if total_percentage >= passing_percentage:
        status = 'Pass'
    else:
        status = 'Fail'
    
    if total_percentage >= 90:
        letter_grade = 'A+'
    elif 80 <= total_percentage < 90:
        letter_grade = 'A'
    elif 70 <= total_percentage < 80:
        letter_grade = 'B'
    elif 60 <= total_percentage < 70:
        letter_grade = 'C'
    elif 50 <= total_percentage < 60:
        letter_grade = 'D'
    else:
        letter_grade = 'F'
        
    student_info['grade'] = letter_grade
    student_info['status'] = status

    return student_info

def get_image_base64(image_bytes):
    """Converts image bytes to a base64 string for embedding in HTML."""
    return base64.b64encode(image_bytes).decode()

def generate_performance_charts(df):
    """Generates performance analysis charts."""
    st.subheader("Performance Analysis")

    # Grade Distribution
    st.write("#### Grade Distribution")
    grade_counts = df['Grade'].value_counts()
    fig, ax = plt.subplots()
    sns.barplot(x=grade_counts.index, y=grade_counts.values, ax=ax)
    ax.set_xlabel("Grade")
    ax.set_ylabel("Number of Students")
    st.pyplot(fig)

    # Subject-wise Performance
    st.write("#### Subject-wise Average Marks")
    
    # Extract subject marks from the 'Total Marks' column
    subject_marks = {}
    for _, row in df.iterrows():
        # The 'Total Marks' column is a string representation of a dictionary
        marks_dict = eval(row['Total Marks'])
        for subject, mark in marks_dict.items():
            if subject not in subject_marks:
                subject_marks[subject] = []
            subject_marks[subject].append(mark)
            
    if subject_marks:
        subject_averages = {subject: sum(marks) / len(marks) for subject, marks in subject_marks.items()}
        
        subject_df = pd.DataFrame(list(subject_averages.items()), columns=['Subject', 'Average Marks'])
        
        fig, ax = plt.subplots()
        sns.barplot(x='Subject', y='Average Marks', data=subject_df, ax=ax)
        ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha="right")
        ax.set_ylabel("Average Marks")
        st.pyplot(fig)

def generate_word_marksheet(student_data, school_name, school_logo_bytes):
    """Generates a well-formatted Word document from a byte stream."""
    document = Document()
    
    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Header
    header_section = document.sections[0].header
    header_paragraph = header_section.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if school_logo_bytes:
        logo_run = header_paragraph.add_run()
        logo_run.add_picture(io.BytesIO(school_logo_bytes), width=Inches(0.8))
        header_paragraph.add_run('\t')
    
    school_name_run = header_paragraph.add_run(school_name)
    school_name_run.font.name = 'Arial'
    school_name_run.font.size = Pt(16)
    school_name_run.bold = True

    # Title
    title = document.add_heading('Student Marksheet', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = 'Arial'
    
    document.add_paragraph() # Spacer

    # Student Profile
    profile_table = document.add_table(rows=1, cols=2)
    profile_table.columns[0].width = Inches(5.0)
    profile_table.columns[1].width = Inches(1.5)
    
    profile_cell = profile_table.cell(0, 0)
    p = profile_cell.paragraphs[0]
    p.add_run('Name: ').bold = True
    p.add_run(f"{student_data['name']}\n")
    p.add_run('Roll No: ').bold = True
    p.add_run(f"{student_data['roll_no']}\n")
    p.add_run('Department: ').bold = True
    p.add_run(f"{student_data['department']}")
    profile_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    photo_cell = profile_table.cell(0, 1)
    photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if student_data.get('photo_bytes'):
        p_photo = photo_cell.paragraphs[0]
        p_photo.add_run().add_picture(io.BytesIO(student_data['photo_bytes']), width=Inches(1.25))
        p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        photo_cell.text = "Photo"
        photo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph() # Spacer

    # Marks table
    document.add_heading('Marks Details', level=2).runs[0].font.name = 'Arial'
    marks_table = document.add_table(rows=1, cols=4)
    marks_table.style = 'Table Grid'
    marks_table.autofit = False
    marks_table.columns[0].width = Inches(2.5)
    marks_table.columns[1].width = Inches(1.5)
    marks_table.columns[2].width = Inches(1.5)
    marks_table.columns[3].width = Inches(1.5)

    hdr_cells = marks_table.rows[0].cells
    headers = ['Subject', 'Mid-term Marks', 'Final Marks', 'Total (Weighted)']
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.add_run(header_text).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for subject, marks in student_data['marks'].items():
        row_cells = marks_table.add_row().cells
        row_cells[0].text = subject
        row_cells[1].text = str(marks['Mid-term'])
        row_cells[2].text = str(marks['Final'])
        row_cells[3].text = f"{student_data['total_marks'].get(subject, 0):.2f}"
        # Center align marks
        for i in range(1, 4):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph() # Spacer

    # Summary
    document.add_heading('Overall Performance', level=2).runs[0].font.name = 'Arial'
    summary_p = document.add_paragraph()
    summary_p.add_run('Final Percentage: ').bold = True
    summary_p.add_run(f"{student_data['percentage']:.2f}%\n")
    summary_p.add_run('GPA: ').bold = True
    summary_p.add_run(f"{student_data['gpa']}\n")
    summary_p.add_run('Grade: ').bold = True
    summary_p.add_run(f"{student_data['grade']}\n")
    summary_p.add_run('Status: ').bold = True
    summary_p.add_run(f"{student_data['status']}")

    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io

def generate_excel_marksheet(student_data, school_name):
    """Generates a well-formatted Excel marksheet for a single student."""
    wb = Workbook()
    ws = wb.active
    ws.title = f"{student_data['name']}"

    # Styles
    center_align = Alignment(horizontal='center', vertical='center')
    right_align = Alignment(horizontal='right', vertical='center')
    bold_font = Font(bold=True, size=12)
    header_font = Font(bold=True, size=16)
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    fail_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")

    # School Header
    ws.merge_cells('A1:D2')
    ws['A1'].value = school_name
    ws['A1'].font = header_font
    ws['A1'].alignment = center_align

    # Student Info
    ws['A4'] = 'Name:'; ws['A4'].font = bold_font
    ws['B4'] = student_data['name']
    ws['A5'] = 'Roll No:'; ws['A5'].font = bold_font
    ws['B5'] = student_data['roll_no']
    ws['A6'] = 'Department:'; ws['A6'].font = bold_font
    ws['B6'] = student_data['department']

    # Marks Table Header
    headers = ['Subject', 'Mid-term Marks', 'Final Marks', 'Total (Weighted)']
    ws.append([]) # Spacer
    header_row = ws.max_row + 1
    for col, text in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col, value=text)
        cell.font = bold_font
        cell.alignment = center_align
        cell.border = thin_border

    # Marks Data
    for subject, marks in student_data['marks'].items():
        data_row = [
            subject,
            marks['Mid-term'],
            marks['Final'],
            f"{student_data['total_marks'].get(subject, 0):.2f}"
        ]
        ws.append(data_row)
        for col in range(1, 5):
            cell = ws.cell(row=ws.max_row, column=col)
            cell.border = thin_border
            if col > 1: cell.alignment = right_align

    # Summary
    summary_row_start = ws.max_row + 2
    ws.cell(row=summary_row_start, column=1, value='Summary').font = bold_font
    
    summary_data = {
        "Percentage:": f"{student_data['percentage']:.2f}%",
        "GPA:": student_data['gpa'],
        "Grade:": student_data['grade'],
        "Status:": student_data['status']
    }
    for i, (key, value) in enumerate(summary_data.items(), 1):
        ws.cell(row=summary_row_start + i, column=1, value=key).font = Font(bold=True)
        cell = ws.cell(row=summary_row_start + i, column=2, value=value)
        if key == "Status:" and value == "Fail":
            cell.fill = fail_fill

    # Adjust Column Widths
    for col_letter in ['A', 'B', 'C', 'D']:
        ws.column_dimensions[col_letter].autosize = True

    excel_io = io.BytesIO()
    wb.save(excel_io)
    excel_io.seek(0)
    return excel_io

# --- Streamlit Dashboard UI ---

st.set_page_config(page_title="Marksheet Input Form", layout="wide")
st.title("Marksheet Generation System")

# Initialize session state
if 'num_subjects' not in st.session_state:
    st.session_state.num_subjects = 3 # Default value
if 'subjects' not in st.session_state:
    st.session_state.subjects = ["Math", "Science", "History"]

def update_subjects():
    st.session_state.subjects = [st.session_state[f"subject_{i}"] for i in range(st.session_state.num_subjects)]

# --- Sidebar for Configuration ---
with st.sidebar:
    st.header("1. School & Academic Criteria")
    school_name = st.text_input("School Name", "Global Tech Academy")
    logo_upload = st.file_uploader("Upload School Logo", type=['png', 'jpg'])
    passing_percentage = st.slider("Passing Percentage (%)", 0, 100, 50)
    mid_term_weight = st.slider("Mid-term Weight", 0.0, 1.0, 0.3, 0.05)
    st.info(f"Final-term Weight: **{1.0 - mid_term_weight:.2f}**")

    st.header("2. Define Subjects")
    num_subjects_input = st.number_input(
        "Number of Subjects", 
        min_value=1, 
        value=st.session_state.num_subjects, 
        step=1,
        key='num_subjects'
    )

    # Dynamically create subject name inputs
    subject_names = []
    for i in range(st.session_state.num_subjects):
        default_name = st.session_state.subjects[i] if i < len(st.session_state.subjects) else f"Subject {i+1}"
        subject_name = st.text_input(f"Subject {i+1} Name", value=default_name, key=f"subject_input_{i}")
        subject_names.append(subject_name)
    
    st.session_state.subjects = subject_names


# --- Main Panel for Input Form ---
st.header("3. Enter Student Details")

if not st.session_state.subjects or not all(st.session_state.subjects):
    st.warning("Please define all subject names in the sidebar to create the input form.")
else:
    with st.form("marksheet_form"):
        st.subheader("Student Information")
        student_name = st.text_input("Student Name", "John Doe")
        roll_no = st.text_input("Roll No", "101")
        department = st.text_input("Department", "Computer Science")
        student_photo = st.file_uploader("Upload Student Photo", type=['png', 'jpg'])

        st.subheader("Enter Marks")
        marks_data = {}
        for subject in st.session_state.subjects:
            cols = st.columns(2)
            with cols[0]:
                mid_marks = st.number_input(f"Mid-term: {subject}", min_value=0, max_value=100, value=75, key=f"mid_{subject}")
            with cols[1]:
                final_marks = st.number_input(f"Final: {subject}", min_value=0, max_value=100, value=80, key=f"final_{subject}")
            marks_data[subject] = {"Mid-term": mid_marks, "Final": final_marks}

        submitted = st.form_submit_button("Generate Marksheet Preview", type="primary")

    if submitted:
        # Process data on submission
        student_info = {
            "name": student_name,
            "roll_no": roll_no,
            "department": department,
            "marks": marks_data,
            "photo_bytes": student_photo.getvalue() if student_photo else None
        }
        
        school_logo_bytes = logo_upload.getvalue() if logo_upload else None

        # Calculate grades
        processed_data = calculate_single_student_grades(student_info, mid_term_weight, (1.0 - mid_term_weight), passing_percentage)

        # --- Display Marksheet Preview ---
        st.header("4. Marksheet Preview & Download")
        
        # Header
        cols = st.columns([1, 4])
        with cols[0]:
            if school_logo_bytes:
                st.image(school_logo_bytes, width=100)
        with cols[1]:
            st.title(school_name)
            st.subheader("Student Marksheet")
        
        st.markdown("---")

        # Student Info
        cols = st.columns([4, 1])
        with cols[0]:
            st.write(f"**Name:** {processed_data['name']}")
            st.write(f"**Roll No:** {processed_data['roll_no']}")
            st.write(f"**Department:** {processed_data['department']}")
        with cols[1]:
            if processed_data['photo_bytes']:
                st.image(processed_data['photo_bytes'], width=120)

        # Marks Table
        st.subheader("Marks Details")
        marks_df_data = []
        for subject, marks in processed_data['marks'].items():
            marks_df_data.append({
                "Subject": subject,
                "Mid-term Marks": marks['Mid-term'],
                "Final Marks": marks['Final'],
                "Total (Weighted)": f"{processed_data['total_marks'][subject]:.2f}"
            })
        st.dataframe(pd.DataFrame(marks_df_data), use_container_width=True)

        # Summary
        st.subheader("Overall Performance")
        cols = st.columns(4)
        with cols[0]:
            st.metric("Final Percentage", f"{processed_data['percentage']:.2f}%")
        with cols[1]:
            st.metric("GPA", processed_data['gpa'])
        with cols[2]:
            st.metric("Grade", processed_data['grade'])
        with cols[3]:
            st.metric("Status", processed_data['status'])

        st.markdown("---")

        # --- Download Buttons ---
        st.subheader("Download Files")
        
        # Generate files in memory
        word_file = generate_word_marksheet(processed_data, school_name, school_logo_bytes)
        excel_file = generate_excel_marksheet(processed_data, school_name)

        cols = st.columns(2)
        with cols[0]:
            st.download_button(
                label="Download as Word (.docx)",
                data=word_file,
                file_name=f"{processed_data['name']}_marksheet.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with cols[1]:
            st.download_button(
                label="Download as Excel (.xlsx)",
                data=excel_file,
                file_name=f"{processed_data['name']}_marksheet.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
